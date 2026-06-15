import docx
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

BASE = Path(__file__).parent

for fname in ['submission_spec.docx', 'redrob_signals_doc.docx', 'README.docx']:
    print(f"\n\n========== {fname} ==========\n")
    try:
        doc = docx.Document(BASE / fname)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        for table in doc.tables:
            print("\n--- TABLE ---")
            for row in table.rows:
                print(" | ".join(cell.text.strip() for cell in row.cells))
    except Exception as e:
        print(f"ERROR reading {fname}: {e}")
